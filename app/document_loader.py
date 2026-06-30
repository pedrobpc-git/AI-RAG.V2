import json
import os
import xml.etree.ElementTree as ET
from html import unescape

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader


def read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as file:
        return file.read()


def read_pdf(path: str) -> str:
    reader = PdfReader(path)
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"Page: {page_number}\n{text}")
    return "\n\n".join(pages)


def read_docx(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    table_text = []
    for table_index, table in enumerate(doc.tables, start=1):
        table_text.append(f"Table: {table_index}")
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                table_text.append(" | ".join(values))

    return "\n".join(paragraphs + table_text)


def _row_to_text(row) -> str:
    values = []
    for column, value in row.items():
        if pd.isna(value):
            continue
        values.append(f"{column}: {value}")
    return " | ".join(values)


def read_csv(path: str) -> str:
    df = pd.read_csv(path, encoding="utf-8", encoding_errors="ignore")
    rows = []
    for index, row in df.iterrows():
        row_text = _row_to_text(row)
        if row_text:
            rows.append(f"Row: {index + 1} | {row_text}")
    return "\n".join(rows)


def read_excel(path: str) -> str:
    sheets = pd.read_excel(path, sheet_name=None)
    parts = []
    for sheet_name, df in sheets.items():
        parts.append(f"Sheet: {sheet_name}")
        for index, row in df.iterrows():
            row_text = _row_to_text(row)
            if row_text:
                parts.append(f"Row: {index + 1} | {row_text}")
    return "\n".join(parts)


def read_json(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as file:
        data = json.load(file)
    return json.dumps(data, indent=2, ensure_ascii=False)


def read_xml(path: str) -> str:
    tree = ET.parse(path)
    root = tree.getroot()
    lines = []

    def walk(node, prefix=""):
        tag = node.tag.split("}")[-1]
        attrs = " ".join(f"{key}={value}" for key, value in node.attrib.items())
        text = (node.text or "").strip()
        line = f"{prefix}{tag}"
        if attrs:
            line += f" ({attrs})"
        if text:
            line += f": {text}"
        lines.append(line)
        for child in list(node):
            walk(child, prefix + "  ")

    walk(root)
    return "\n".join(lines)


def read_html(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as file:
        soup = BeautifulSoup(file.read(), "html.parser")
    for element in soup(["script", "style", "noscript"]):
        element.decompose()
    return unescape(soup.get_text(separator="\n", strip=True))


def load_documents(directory: str) -> list[dict]:
    documents = []

    if not os.path.isdir(directory):
        return documents

    for root, _, files in os.walk(directory):
        for filename in files:
            path = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()

            try:
                if ext in {".txt", ".md"}:
                    text = read_txt(path)
                elif ext == ".pdf":
                    text = read_pdf(path)
                elif ext == ".docx":
                    text = read_docx(path)
                elif ext == ".csv":
                    text = read_csv(path)
                elif ext in {".xls", ".xlsx", ".xlsm"}:
                    text = read_excel(path)
                elif ext == ".json":
                    text = read_json(path)
                elif ext == ".xml":
                    text = read_xml(path)
                elif ext in {".html", ".htm"}:
                    text = read_html(path)
                else:
                    continue
            except Exception as exc:
                print(f"[WARN] Nu pot citi {path}: {exc}")
                continue

            if text and text.strip():
                documents.append({"source": path, "text": text})

    return documents
